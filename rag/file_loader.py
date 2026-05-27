import os
import uuid
from PIL import Image
import pytesseract
 
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from docx import Document as DocxDocument
 
def load_files(files):
    docs = []
 
    # Create temp directory if it doesn't exist
    os.makedirs("temp", exist_ok=True) 
 
    for f in files:
        # Generate a unique filename to eliminate structural collisions or overwrites
        unique_name = f"{uuid.uuid4()}_{f.name}"
        path = f"temp/{unique_name}"
 
        # Write binary stream cache down to temporary space
        with open(path, "wb") as t:
            t.write(f.getbuffer())
 
        try:
            # ✅ Native Vector extraction for PDF strings
            if f.name.endswith(".pdf"):
                docs.extend(PyPDFLoader(path).load())
 
            # ✅ Plain text file loader
            elif f.name.endswith(".txt"):
                docs.extend(TextLoader(path).load())
 
            # ✅ Microsoft Open XML Document parser
            elif f.name.endswith(".docx"):
                doc = DocxDocument(path)
                text = "\n".join([para.text for para in doc.paragraphs])
                docs.append(Document(page_content=text, metadata={"source": f.name}))
 
            # ✅ Optical Character Recognition (OCR) Engine pipeline for Images
            elif f.name.lower().endswith((".png", ".jpg", ".jpeg")):
                text = pytesseract.image_to_string(Image.open(path))
                docs.append(Document(page_content=text, metadata={"source": f.name}))
 
            else:
                print(f"Unsupported format array verification failure for: {f.name}")
                continue
 
        except Exception as e:
            print(f"Critical operational error parsing file trace context {f.name}: {e}")
        finally:
            # Clean up the individual file from disk immediately after processing to keep local environment light
            if os.path.exists(path):
                os.remove(path)
 
    return docs